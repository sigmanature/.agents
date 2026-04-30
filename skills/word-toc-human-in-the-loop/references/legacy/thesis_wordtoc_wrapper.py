#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
from pathlib import Path

import thesis as T

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent


def _apply_heading_styles(doc: Document):
    """把 Heading 1/2/3 调成你原来 add_heading 的视觉效果，但让 Word 能识别为标题层级。"""
    h1 = doc.styles["Heading 1"]
    h1.font.name = "宋体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    h1.font.size = Pt(16)
    h1.font.bold = False
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(0)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "宋体"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(0)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "宋体"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    h3.font.size = Pt(12)
    h3.font.bold = False
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(0)
    h3.paragraph_format.space_after = Pt(0)


def _add_toc_field(doc: Document, levels: str = "1-3"):
    """插入 Word TOC 域：页码由 Word 计算。"""
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f'TOC \\o "{levels}" \\h \\z \\u')
    p._p.append(fld)
    return p


def _placeholder_pic(doc: Document, msg: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"【图片缺失：{msg}】")
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 0, 0)


def build(out_docx: str = "Final_Thesis_WordTOC.docx"):
    # ---- 1) 先保存 thesis.py 里原函数引用，避免递归 ----
    _orig_add_heading = T.add_heading
    _orig_add_figure = T.add_figure

    # ---- 2) monkey patch：标题 level 1/2/3 用 Heading 1/2/3 ----
    def patched_add_heading(doc: Document, text: str, level: int):
        if level in (1, 2, 3):
            p = doc.add_paragraph(style=f"Heading {level}")
            run = p.add_run(text)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            # 保险：显式设一下，防止模板改过样式
            if level == 1:
                run.font.size = Pt(16); run.bold = False
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                run.font.size = Pt(14); run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                run.font.size = Pt(12); run.bold = False
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            return p
        return _orig_add_heading(doc, text, level)

    # ---- 3) monkey patch：图片路径相对化 + 缺失占位；关键：调用 _orig_add_figure ----
    def patched_add_figure(doc: Document, image_path: str, caption_text: str,
                           chapter_no=None, width_cm=None, scale: float = 0.9):
        p = Path(image_path)
        if not p.is_absolute():
            cand = BASE_DIR / image_path
            if cand.exists():
                image_path = str(cand)

        if not Path(image_path).exists():
            _placeholder_pic(doc, f"{image_path}（{caption_text}）")
            # 仍然给个图题（不影响分页，只是让你知道缺啥）
            return None

        return _orig_add_figure(doc, image_path, caption_text,
                                chapter_no=chapter_no, width_cm=width_cm, scale=scale)

    T.add_heading = patched_add_heading
    T.add_figure = patched_add_figure

    # ---- 4) 生成文档 ----
    doc = Document()
    T.setup_page(doc)
    _apply_heading_styles(doc)

    # 目录标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("目  录")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(16)
    r.bold = False

    doc.add_paragraph()
    _add_toc_field(doc, levels="1-3")

    # 分节：正文从 1 开始（目录不算正文页码）
    doc.add_section(WD_SECTION.NEW_PAGE)
    sectPr = doc.sections[1]._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), "1")

    # 正文（复用你 thesis.py 原逻辑）
    T.generate_chapter_1(doc)
    T.generate_chapter_2(doc)
    T.generate_chapter_3(doc)
    T.generate_chapter_4(doc)
    T.generate_chapter_5(doc)

    # 你后续要的话：这里可以继续追加 参考文献/致谢（也会进目录，因为是 Heading 1）
    doc.save(out_docx)
    print(f"[OK] generated: {out_docx}")


if __name__ == "__main__":
    build()
