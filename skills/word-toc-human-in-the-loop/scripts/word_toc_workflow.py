#!/usr/bin/env python3

from __future__ import annotations

import argparse
import math
import os
import re
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile

from PIL import ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
TOC_LEVEL_1_LABELS = {"摘要", "ABSTRACT", "结语", "参考文献", "致谢"}
FONT_CANDIDATES = {
    "song": [
        "/usr/share/fonts/truetype/arphic/uming.ttc",
        "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSerifCJK-VF.ttc",
    ],
    "hei": [
        "/usr/share/fonts/truetype/arphic/ukai.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-VF.ttc",
    ],
}


@dataclass
class HeadingItem:
    level: int
    text: str
    index: int


@dataclass
class TocItem:
    level: int
    text: str
    page: str
    index: int


@dataclass
class DocSnapshot:
    heading_items: list[HeadingItem]
    toc_items: list[TocItem]
    toc_title_index: int
    first_toc_entry_index: int
    last_toc_entry_index: int
    first_body_heading_index: int


@dataclass
class AuditReport:
    is_dirty: bool
    missing_in_toc: list[str]
    extra_in_toc: list[str]
    page_numbering_ok: bool


@dataclass
class TocTemplate:
    page_width_pt: float
    section_left_margin_pt: float
    section_right_margin_pt: float
    section_top_margin_pt: float
    section_bottom_margin_pt: float
    title_text: str
    title_font_name: str
    title_font_size_pt: float
    title_line_spacing_pt: float
    level1_alignment: int
    level1_left_pt: float | None
    level1_first_line_pt: float | None
    level1_title_font_name: str
    level1_title_font_size_pt: float
    level2_alignment: int
    level2_left_pt: float | None
    level2_first_line_pt: float | None
    level2_title_font_name: str
    level2_title_font_size_pt: float
    level3_alignment: int
    level3_left_pt: float | None
    level3_first_line_pt: float | None
    level3_title_font_name: str
    level3_title_font_size_pt: float
    dots_font_name: str
    dots_font_size_pt: float


def _normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _normalize_special_toc_title(text: str) -> str:
    compact = text.replace(" ", "").strip()
    if compact == "摘要":
        return "摘要"
    if text.strip().lower() == "abstract":
        return "ABSTRACT"
    return _normalize_spaces(text)


def _pt_or_none(length) -> float | None:
    return length.pt if length is not None else None


def _env_font_candidates(kind: str) -> list[str]:
    env_name = f"WORD_TOC_{kind.upper()}_FONTS"
    raw = os.environ.get(env_name, "").strip()
    if not raw:
        return []
    return [item for item in raw.split(os.pathsep) if item]


def _candidate_paths(kind: str) -> list[str]:
    env_candidates = _env_font_candidates(kind)
    default_candidates = FONT_CANDIDATES[kind]
    return env_candidates + default_candidates


def _load_font(kind: str, size_pt: float) -> ImageFont.FreeTypeFont:
    for candidate in _candidate_paths(kind):
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), round(size_pt))
    env_name = f"WORD_TOC_{kind.upper()}_FONTS"
    raise RuntimeError(
        f"font candidate for {kind} not found; set {env_name} to one or more font files separated by '{os.pathsep}'"
    )


def _font_kind(font_name: str) -> str:
    return "hei" if "黑" in font_name else "song"


def _set_run_font(run, font_name: str, size_pt: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = False
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def _delete_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def _style_to_level(style_name: str) -> int | None:
    mapping = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3}
    return mapping.get(style_name)


def _outline_to_level(paragraph) -> int | None:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return None
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        return None
    value = outline.get(qn("w:val"))
    if value is None:
        return None
    try:
        outline = int(value)
    except ValueError:
        return None
    if 0 <= outline <= 2:
        return outline + 1
    return None


def _find_toc_title_index(doc: Document) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().replace(" ", "") == "目录":
            return idx
    raise RuntimeError("TOC title '目 录' not found")


def _find_first_body_heading_index(doc: Document, start_index: int) -> int:
    for idx in range(start_index + 1, len(doc.paragraphs)):
        if doc.paragraphs[idx].style.name == "Heading 1" and doc.paragraphs[idx].text.strip():
            return idx
    raise RuntimeError("first body Heading 1 after TOC not found")


def _infer_toc_level(title: str) -> int:
    if title in TOC_LEVEL_1_LABELS or re.match(r"^第.+章", title):
        return 1
    if re.match(r"^\d+\.\d+\.\d+", title):
        return 3
    if re.match(r"^\d+\.\d+", title):
        return 2
    return 1


def _parse_toc_line(text: str) -> tuple[str, str] | None:
    raw = text.strip()
    if not raw:
        return None

    if "\t" in raw:
        left, page = raw.rsplit("\t", 1)
        title = _normalize_special_toc_title(left.replace("…", ""))
        page = page.strip()
        if title and page:
            return title, page

    match = re.match(r"^(.*?)([IVXLC]+|\d+)$", raw)
    if not match:
        return None

    title = _normalize_special_toc_title(match.group(1).replace("…", ""))
    page = match.group(2)
    if not title:
        return None
    return title, page


def extract_heading_items(doc: Document) -> list[HeadingItem]:
    items: list[HeadingItem] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        level = _style_to_level(paragraph.style.name)
        if level is None:
            level = _outline_to_level(paragraph)
        if level is None:
            continue
        text = _normalize_special_toc_title(paragraph.text)
        if text:
            items.append(HeadingItem(level=level, text=text, index=idx))
    return items


def extract_doc_snapshot(docx_path: Path | str) -> DocSnapshot:
    path = Path(docx_path)
    doc = Document(path)

    toc_title_index = _find_toc_title_index(doc)
    first_body_heading_index = _find_first_body_heading_index(doc, toc_title_index)

    toc_items: list[TocItem] = []
    for idx in range(toc_title_index + 1, first_body_heading_index):
        parsed = _parse_toc_line(doc.paragraphs[idx].text)
        if parsed is None:
            continue
        title, page = parsed
        toc_items.append(TocItem(level=_infer_toc_level(title), text=title, page=page, index=idx))

    if not toc_items:
        raise RuntimeError("no TOC entries found")

    return DocSnapshot(
        heading_items=extract_heading_items(doc),
        toc_items=toc_items,
        toc_title_index=toc_title_index,
        first_toc_entry_index=toc_items[0].index,
        last_toc_entry_index=toc_items[-1].index,
        first_body_heading_index=first_body_heading_index,
    )


def extract_toc_template(docx_path: Path | str) -> TocTemplate:
    path = Path(docx_path)
    doc = Document(path)
    snapshot = extract_doc_snapshot(path)
    section = doc.sections[0]

    def _level_sample(level: int):
        item = next(item for item in snapshot.toc_items if item.level == level)
        paragraph = doc.paragraphs[item.index]
        title_run = paragraph.runs[0]
        return item, paragraph, title_run

    title_para = doc.paragraphs[snapshot.toc_title_index]
    _, level1_para, level1_run = _level_sample(1)
    _, level2_para, level2_run = _level_sample(2)
    _, level3_para, level3_run = _level_sample(3)
    dots_run = next(
        (run for run in level3_para.runs[1:] if run.font.name or run.font.size),
        level3_run,
    )

    return TocTemplate(
        page_width_pt=section.page_width.pt,
        section_left_margin_pt=section.left_margin.pt,
        section_right_margin_pt=section.right_margin.pt,
        section_top_margin_pt=section.top_margin.pt,
        section_bottom_margin_pt=section.bottom_margin.pt,
        title_text=title_para.text,
        title_font_name=title_para.runs[0].font.name,
        title_font_size_pt=title_para.runs[0].font.size.pt,
        title_line_spacing_pt=title_para.paragraph_format.line_spacing.pt,
        level1_alignment=level1_para.alignment,
        level1_left_pt=_pt_or_none(level1_para.paragraph_format.left_indent),
        level1_first_line_pt=_pt_or_none(level1_para.paragraph_format.first_line_indent),
        level1_title_font_name=level1_run.font.name,
        level1_title_font_size_pt=level1_run.font.size.pt,
        level2_alignment=level2_para.alignment,
        level2_left_pt=_pt_or_none(level2_para.paragraph_format.left_indent),
        level2_first_line_pt=_pt_or_none(level2_para.paragraph_format.first_line_indent),
        level2_title_font_name=level2_run.font.name,
        level2_title_font_size_pt=level2_run.font.size.pt,
        level3_alignment=level3_para.alignment,
        level3_left_pt=_pt_or_none(level3_para.paragraph_format.left_indent),
        level3_first_line_pt=_pt_or_none(level3_para.paragraph_format.first_line_indent),
        level3_title_font_name=level3_run.font.name,
        level3_title_font_size_pt=level3_run.font.size.pt,
        dots_font_name=dots_run.font.name,
        dots_font_size_pt=dots_run.font.size.pt,
    )


def has_body_page_number_restart(docx_path: Path | str) -> bool:
    path = Path(docx_path)
    with ZipFile(path) as package:
        root = etree.fromstring(package.read("word/document.xml"))

    sect_prs = root.xpath("//w:sectPr", namespaces=NS)
    for sect_pr in sect_prs[1:]:
        starts = sect_pr.xpath("./w:pgNumType/@w:start", namespaces=NS)
        if starts and starts[0] == "1":
            return True
    return False


def _ensure_body_page_number_restart(doc: Document) -> None:
    if len(doc.sections) < 2:
        return
    sect_pr = doc.sections[1]._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), "1")


def _ensure_outline_level(paragraph, level: int) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level - 1))


def _insert_toc_field(paragraph, levels: str = "1-3") -> None:
    _clear_paragraph(paragraph)
    paragraph.style = "Normal"
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f'TOC \\o "{levels}" \\h \\z \\u')
    paragraph._p.append(fld)


def prepare_word_toc_docx(input_docx: Path | str, output_docx: Path | str, levels: str = "1-3") -> None:
    path = Path(input_docx)
    doc = Document(path)
    snapshot = extract_doc_snapshot(path)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        compact = text.replace(" ", "")
        if compact == "摘要" or text.lower() == "abstract":
            _ensure_outline_level(paragraph, 1)

    _insert_toc_field(doc.paragraphs[snapshot.first_toc_entry_index], levels=levels)
    for idx in range(snapshot.last_toc_entry_index, snapshot.first_toc_entry_index, -1):
        _delete_paragraph(doc.paragraphs[idx])

    _ensure_body_page_number_restart(doc)
    Path(output_docx).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def audit_docx(docx_path: Path | str) -> AuditReport:
    snapshot = extract_doc_snapshot(docx_path)
    heading_texts = [item.text for item in snapshot.heading_items if item.text not in {"摘要", "ABSTRACT"}]
    toc_texts = [item.text for item in snapshot.toc_items if item.text not in {"摘要", "ABSTRACT"}]

    missing_in_toc = [text for text in heading_texts if text not in toc_texts]
    extra_in_toc = [text for text in toc_texts if text not in heading_texts]

    return AuditReport(
        is_dirty=bool(missing_in_toc or extra_in_toc),
        missing_in_toc=missing_in_toc,
        extra_in_toc=extra_in_toc,
        page_numbering_ok=has_body_page_number_restart(docx_path),
    )


def _toc_level_from_paragraph(paragraph, parsed_title: str) -> int:
    style_name = paragraph.style.name.lower()
    if style_name == "toc 1":
        return 1
    if style_name == "toc 2":
        return 2
    if style_name == "toc 3":
        return 3
    return _infer_toc_level(parsed_title)


def _format_toc_title(paragraph, template: TocTemplate) -> None:
    _clear_paragraph(paragraph)
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.right_indent = None
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(template.title_line_spacing_pt)
    run = paragraph.add_run(template.title_text)
    _set_run_font(run, template.title_font_name, template.title_font_size_pt)


def _format_blank(paragraph) -> None:
    _clear_paragraph(paragraph)
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.right_indent = None
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run("")


def _ellipsis_text(title: str, page: str, template: TocTemplate, level: int) -> str:
    if level == 1:
        left_pt = (template.level1_left_pt or 0.0) + (template.level1_first_line_pt or 0.0)
        font_name = template.level1_title_font_name
        font_size = template.level1_title_font_size_pt
    elif level == 2:
        left_pt = (template.level2_left_pt or 0.0) + (template.level2_first_line_pt or 0.0)
        font_name = template.level2_title_font_name
        font_size = template.level2_title_font_size_pt
    else:
        left_pt = (template.level3_left_pt or 0.0) + (template.level3_first_line_pt or 0.0)
        font_name = template.level3_title_font_name
        font_size = template.level3_title_font_size_pt

    title_font = _load_font(_font_kind(font_name), font_size)
    dots_font = _load_font("song", template.dots_font_size_pt)
    title_width = title_font.getlength(title)
    page_width = dots_font.getlength(page)
    ellipsis_width = dots_font.getlength("…")
    content_width = template.page_width_pt - template.section_left_margin_pt - template.section_right_margin_pt
    available = content_width - left_pt
    count = max(2, math.floor(max(0.0, available - title_width - page_width) / max(1.0, ellipsis_width)))
    return "…" * count


def _apply_level_format(paragraph, template: TocTemplate, level: int) -> tuple[str, float]:
    paragraph.style = "Normal"
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.right_indent = None

    if level == 1:
        paragraph.alignment = template.level1_alignment
        paragraph.paragraph_format.left_indent = Pt(template.level1_left_pt) if template.level1_left_pt is not None else None
        paragraph.paragraph_format.first_line_indent = (
            Pt(template.level1_first_line_pt) if template.level1_first_line_pt is not None else None
        )
        return template.level1_title_font_name, template.level1_title_font_size_pt

    if level == 2:
        paragraph.alignment = template.level2_alignment
        paragraph.paragraph_format.left_indent = Pt(template.level2_left_pt) if template.level2_left_pt is not None else None
        paragraph.paragraph_format.first_line_indent = (
            Pt(template.level2_first_line_pt) if template.level2_first_line_pt is not None else None
        )
        return template.level2_title_font_name, template.level2_title_font_size_pt

    paragraph.alignment = template.level3_alignment
    paragraph.paragraph_format.left_indent = Pt(template.level3_left_pt) if template.level3_left_pt is not None else None
    paragraph.paragraph_format.first_line_indent = (
        Pt(template.level3_first_line_pt) if template.level3_first_line_pt is not None else None
    )
    return template.level3_title_font_name, template.level3_title_font_size_pt


def restyle_docx(input_docx: Path | str, template_docx: Path | str, output_docx: Path | str) -> None:
    template = extract_toc_template(template_docx)
    path = Path(input_docx)
    doc = Document(path)

    section = doc.sections[0]
    section.left_margin = Pt(template.section_left_margin_pt)
    section.right_margin = Pt(template.section_right_margin_pt)
    section.top_margin = Pt(template.section_top_margin_pt)
    section.bottom_margin = Pt(template.section_bottom_margin_pt)

    toc_title_index = _find_toc_title_index(doc)
    first_body_heading_index = _find_first_body_heading_index(doc, toc_title_index)

    _format_toc_title(doc.paragraphs[toc_title_index], template)
    if toc_title_index + 1 < len(doc.paragraphs) and _parse_toc_line(doc.paragraphs[toc_title_index + 1].text) is None:
        _format_blank(doc.paragraphs[toc_title_index + 1])

    formatted_count = 0
    for idx in range(toc_title_index + 1, first_body_heading_index):
        paragraph = doc.paragraphs[idx]
        parsed = _parse_toc_line(paragraph.text)
        if parsed is None:
            continue
        title, page = parsed
        level = _toc_level_from_paragraph(paragraph, title)
        title_font_name, title_font_size = _apply_level_format(paragraph, template, level)
        dots = _ellipsis_text(title, page, template, level)
        _clear_paragraph(paragraph)
        title_run = paragraph.add_run(title)
        _set_run_font(title_run, title_font_name, title_font_size)
        dots_run = paragraph.add_run(dots)
        _set_run_font(dots_run, template.dots_font_name, template.dots_font_size_pt)
        page_run = paragraph.add_run(page)
        _set_run_font(page_run, template.dots_font_name, template.dots_font_size_pt)
        formatted_count += 1

    if formatted_count == 0:
        raise RuntimeError("no TOC entries found to restyle; update the TOC in Word first")

    Path(output_docx).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser()
    subparsers = parser.add_subparsers(dest="command", required=True)

    subparsers.add_parser("doctor")

    audit = subparsers.add_parser("audit")
    audit.add_argument("input_docx", type=Path)

    prepare = subparsers.add_parser("prepare")
    prepare.add_argument("input_docx", type=Path)
    prepare.add_argument("--output", required=True, type=Path)
    prepare.add_argument("--levels", default="1-3")

    restyle = subparsers.add_parser("restyle")
    restyle.add_argument("input_docx", type=Path)
    restyle.add_argument("--template", required=True, type=Path)
    restyle.add_argument("--output", required=True, type=Path)

    return parser


def _doctor() -> int:
    print("font_candidates:")
    for kind in ("song", "hei"):
        print(f"  {kind}:")
        for candidate in _candidate_paths(kind):
            status = "ok" if Path(candidate).exists() else "missing"
            print(f"    - [{status}] {candidate}")
    print("env_override_help:")
    print(f"  WORD_TOC_SONG_FONTS: paths separated by '{os.pathsep}'")
    print(f"  WORD_TOC_HEI_FONTS: paths separated by '{os.pathsep}'")
    return 0


def main(argv: list[str] | None = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)

    if args.command == "doctor":
        return _doctor()

    if args.command == "audit":
        report = audit_docx(args.input_docx)
        print(f"is_dirty: {report.is_dirty}")
        print(f"page_numbering_ok: {report.page_numbering_ok}")
        print("missing_in_toc:")
        for item in report.missing_in_toc:
            print(f"  - {item}")
        print("extra_in_toc:")
        for item in report.extra_in_toc:
            print(f"  - {item}")
        return 0

    if args.command == "prepare":
        prepare_word_toc_docx(args.input_docx, args.output, levels=args.levels)
        print(f"prepared: {args.output}")
        return 0

    if args.command == "restyle":
        restyle_docx(args.input_docx, args.template, args.output)
        print(f"restyled: {args.output}")
        return 0

    parser.error(f"unsupported command: {args.command}")
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
